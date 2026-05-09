import os
import tempfile

import requests
from docx import Document


def main():
    tmp = tempfile.gettempdir()
    file_path = os.path.join(tmp, "test_upload.docx")

    doc = Document()
    doc.add_heading("装备系统", level=1)
    doc.add_paragraph("这是一个测试文档，用于验证上传。")
    doc.save(file_path)

    endpoints = [
        ("文档库上传", "http://127.0.0.1:8000/api/documents/upload"),
        ("知识库入库", "http://127.0.0.1:8000/api/kb/upload"),
    ]

    for name, url in endpoints:
        with open(file_path, "rb") as f:
            files = {
                "file": (
                    "装备系统.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            data = {"doc_type": "system"}
            r = requests.post(url, files=files, data=data, timeout=120)
            print("=", name, "=")
            print("status", r.status_code)
            print(r.text[:800])


if __name__ == "__main__":
    main()

