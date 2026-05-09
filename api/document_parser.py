import os
import base64
import io
import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# 将图片转成 base64 data URI
def _image_to_data_uri(image_part) -> str:
    try:
        content_type = image_part.content_type
        if not content_type or not content_type.startswith("image/"):
            content_type = "image/png"
        raw = image_part.blob
        b64 = base64.b64encode(raw).decode("ascii")
        return f"data:{content_type};base64,{b64}"
    except Exception:
        return ""


def _parse_table(table) -> str:
    """将 python-docx 的 Table 对象转成 HTML <table>"""
    rows_html = []
    for row in table.rows:
        cells_html = []
        for cell in row.cells:
            tag = "th" if row == table.rows[0] else "td"
            # 递归处理单元格内的段落和内嵌图片
            cell_content = _parse_paragraphs(cell.paragraphs, cell)
            cells_html.append(f"<{tag}>{cell_content}</{tag}>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return '<table style="border-collapse:collapse;width:100%">' + "".join(rows_html) + "</table>"


def _extract_images_from_paragraph(para) -> str:
    """提取段落中内嵌的图片，返回 <img> 标签拼接"""
    parts = []
    for run in para.runs:
        # 检查 run 中是否包含图片
        for child in run._element:
            if child.tag.endswith("drawing"):
                # Inline drawing: 提取 rId → image part
                blips = child.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
                for blip in blips:
                    embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if embed:
                        try:
                            image_part = para.part.related_parts[embed]
                            uri = _image_to_data_uri(image_part)
                            if uri:
                                parts.append(f'<img src="{uri}" alt="图片" style="max-width:100%;height:auto;margin:8px 0;border-radius:4px;display:block" />')
                        except Exception:
                            pass
            elif child.tag.endswith("pict"):
                # VML 图片（旧格式）
                imgs = child.findall(".//{urn:schemas-microsoft-com:vml}imagedata")
                for img in imgs:
                    rid = img.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                    if rid:
                        try:
                            image_part = para.part.related_parts[rid]
                            uri = _image_to_data_uri(image_part)
                            if uri:
                                parts.append(f'<img src="{uri}" alt="图片" style="max-width:100%;height:auto;margin:8px 0;border-radius:4px;display:block" />')
                        except Exception:
                            pass
    return "".join(parts)


def _parse_paragraphs(paragraphs, parent=None) -> str:
    """将段落列表转成 HTML"""
    html_parts = []
    for para in paragraphs:
        img_html = _extract_images_from_paragraph(para)
        if img_html:
            html_parts.append(img_html)
        text = para.text.strip()
        if text:
            # 根据段落样式判断是否为标题
            style_name = (para.style.name or "").lower() if para.style else ""
            if style_name.startswith("heading 1") or style_name.startswith("heading1"):
                html_parts.append(f"<h1>{text}</h1>")
            elif style_name.startswith("heading 2") or style_name.startswith("heading2"):
                html_parts.append(f"<h2>{text}</h2>")
            elif style_name.startswith("heading 3") or style_name.startswith("heading3"):
                html_parts.append(f"<h3>{text}</h3>")
            elif img_html:
                html_parts.append(f"<p>{text}</p>")
            else:
                html_parts.append(f"<p>{text}</p>")
    return "".join(html_parts)


class DocumentParser:
    """负责解析不同格式的文档，输出 HTML（含图片 base64 和表格）"""

    @staticmethod
    def parse_docx(file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            html_parts = []

            for block in doc.element.body:
                tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

                if tag == "p":
                    # 段落：在 doc.paragraphs 中找到对应对象
                    # 简化处理：直接遍历所有顶级段落和表格
                    pass

            # 遍历 body 元素，保持顺序
            body = doc.element.body
            para_index = 0
            table_index = 0
            for child in body:
                ctag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if ctag == "p":
                    if para_index < len(doc.paragraphs):
                        para = doc.paragraphs[para_index]
                        img_html = _extract_images_from_paragraph(para)
                        text = para.text.strip()
                        style_name = (para.style.name or "").lower() if para.style else ""
                        if img_html:
                            html_parts.append(img_html)
                        if text:
                            if style_name.startswith("heading 1") or style_name.startswith("heading1"):
                                html_parts.append(f"<h1>{text}</h1>")
                            elif style_name.startswith("heading 2") or style_name.startswith("heading2"):
                                html_parts.append(f"<h2>{text}</h2>")
                            elif style_name.startswith("heading 3") or style_name.startswith("heading3"):
                                html_parts.append(f"<h3>{text}</h3>")
                            else:
                                html_parts.append(f"<p>{text}</p>")
                        elif not img_html and not text:
                            html_parts.append("<p><br></p>")
                        para_index += 1
                elif ctag == "tbl":
                    if table_index < len(doc.tables):
                        html_parts.append(_parse_table(doc.tables[table_index]))
                        table_index += 1

            result = "".join(html_parts)
            if not result.strip():
                # 兜底：只用段落文本
                result = "".join(f"<p>{p.text}</p>" for p in doc.paragraphs if p.text.strip())
            return result
        except Exception as e:
            # 兜底解析
            try:
                doc = docx.Document(file_path)
                return "".join(f"<p>{p.text}</p>" for p in doc.paragraphs if p.text.strip())
            except Exception:
                return f"<p>DOCX 解析错误: {str(e)}</p>"

    @staticmethod
    def parse_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return "<p>" + text.replace("\n\n", "</p><p>").replace("\n", "<br>") + "</p>"

    @staticmethod
    def parse_md(file_path: str) -> str:
        try:
            import markdown
            with open(file_path, 'r', encoding='utf-8') as f:
                return markdown.markdown(f.read(), extensions=['tables', 'fenced_code'])
        except ImportError:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return "<pre>" + text + "</pre>"

    @classmethod
    def parse(cls, file_path: str) -> str:
        file_path = str(file_path)
        if not os.path.exists(file_path):
            return "<p>文件不存在</p>"
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.docx':
            return cls.parse_docx(file_path)
        elif ext == '.txt':
            return cls.parse_txt(file_path)
        elif ext == '.md':
            return cls.parse_md(file_path)
        elif ext in ('.xlsx', '.xls'):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, data_only=True)
                ws = wb.active
                if not ws:
                    return "<p>空表格</p>"
                html = ['<table style="border-collapse:collapse;width:100%">']
                for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 500), values_only=True):
                    html.append('<tr>' + ''.join(f'<td style="border:1px solid #d4d4d8;padding:6px 10px">{v if v is not None else ""}</td>' for v in row) + '</tr>')
                html.append('</table>')
                wb.close()
                return ''.join(html)
            except ImportError:
                return "<p>无法解析 Excel 文件（缺少 openpyxl）</p>"
            except Exception as e:
                return f"<p>Excel 解析错误: {str(e)}</p>"
        else:
            return "<p>不支持的文件格式</p>"
