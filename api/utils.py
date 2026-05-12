import os
import io
import json
import time
import uuid
import base64
import traceback
import tempfile
import re


def get_app_data_dir():
    override_dir = os.environ.get('GB_DATA_DIR')
    if override_dir:
        os.makedirs(override_dir, exist_ok=True)
        return override_dir

    if os.name == 'nt':
        app_data = os.environ.get('APPDATA')
        if not app_data:
            app_data = os.path.join(os.path.expanduser('~'), 'AppData', 'Roaming')
        target_dir = os.path.join(app_data, 'GameBuilderAIHelper')
    else:
        target_dir = os.path.join(os.path.expanduser('~'), '.gamebuilder_aihelper')

    os.makedirs(target_dir, exist_ok=True)
    return target_dir


def ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)


def load_json(filepath, default=None):
    if default is None:
        default = {}
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return default


def save_json(filepath, data):
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def log_error(title: str, err: Exception, log_dir: str):
    try:
        log_file = os.path.join(log_dir, "api.log")
        with open(log_file, 'a', encoding='utf-8') as f:
            f.write('\n' + '=' * 80 + '\n')
            f.write(f"{int(time.time())} {title}\n")
            f.write(str(err) + '\n')
            f.write(traceback.format_exc() + '\n')
    except Exception:
        pass


def save_docx_to_path(content: str, name: str, output_path: str) -> str:
    """将 HTML 内容生成为 .docx 并保存到指定路径，返回文件路径。"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    clean_name = name.replace('.docx', '').replace('.doc', '').strip()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    if clean_name:
        h = doc.add_heading(clean_name, level=1)
        for run in h.runs:
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    _html_to_docx(content or '', doc)

    if not clean_name and not (content or '').strip():
        doc.add_paragraph('（空文档）')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_docx(content: str, name: str) -> dict:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from html.parser import HTMLParser
    import struct

    doc = Document()
    clean_name = name.replace('.docx', '').replace('.doc', '').strip()

    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    if clean_name:
        h = doc.add_heading(clean_name, level=1)
        for run in h.runs:
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Parse HTML and build docx content
    _html_to_docx(content or '', doc)

    if not clean_name and not (content or '').strip():
        doc.add_paragraph('（空文档）')

    tmp_path = os.path.join(tempfile.gettempdir(), f'gb_docx_{uuid.uuid4().hex}.docx')
    doc.save(tmp_path)
    with open(tmp_path, 'rb') as f:
        b64 = base64.b64encode(f.read()).decode('ascii')
    try:
        os.unlink(tmp_path)
    except Exception:
        pass
    return {"success": True, "data_uri": f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}"}


def _html_to_docx(html: str, doc):
    """Convert Tiptap HTML to python-docx formatted document content."""
    from html.parser import HTMLParser
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    # Decode HTML entities
    html = html.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    html = html.replace('&nbsp;', ' ').replace('&quot;', '"').replace('&#39;', "'")

    class DocxHTMLParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.doc = doc
            self.para = None
            self.table = None
            self.tr = None
            self.td = None
            self.in_table = False
            self.in_td = False
            self.in_li = False
            self.list_level = 0
            self.list_type = None  # 'ul' or 'ol'
            self.list_counter = [0]
            self.skip_next_data = False
            self.text_buffer = []

            # Formatting stack
            self.runs_stack = []  # stack of formatting dicts
            self.current_formats = {
                'bold': False, 'italic': False, 'underline': False,
                'strikethrough': False, 'code': False,
                'color': None, 'highlight': None, 'font_size': None,
                'font_family': None, 'link_url': None,
            }

        def _add_run(self, text, formats=None):
            if not self.para:
                self.para = self.doc.add_paragraph()
            fmt = formats or self.current_formats
            run = self.para.add_run(text)
            if fmt['bold']:
                run.bold = True
            if fmt['italic']:
                run.italic = True
            if fmt['underline']:
                run.underline = True
            if fmt['strikethrough']:
                run.font.strike = True
            if fmt['code']:
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            if fmt['color']:
                try:
                    run.font.color.rgb = RGBColor.from_string(fmt['color'].lstrip('#'))
                except Exception:
                    pass
            if fmt['highlight']:
                try:
                    from docx.oxml import OxmlElement
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), fmt['highlight'].lstrip('#'))
                    shd.set(qn('w:val'), 'clear')
                    run.element.rPr.append(shd)
                except Exception:
                    pass
            if fmt['font_size']:
                try:
                    run.font.size = Pt(float(fmt['font_size'].replace('px', '')))
                except Exception:
                    pass
            if fmt['font_family']:
                run.font.name = fmt['font_family']
                run.element.rPr.rFonts.set(qn('w:eastAsia'), fmt['font_family'])
            if fmt['link_url']:
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                try:
                    rel = self.para.part.relate_to(fmt['link_url'], RT.HYPERLINK, is_external=True)
                    from docx.oxml import OxmlElement
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), rel)
                    r_elem = run.element
                    r_parent = r_elem.getparent()
                    hyperlink.append(r_elem)
                    r_parent.replace(r_elem, hyperlink)
                except Exception:
                    pass

        def _add_image(self, data_uri):
            """Add an inline image from a data URI."""
            if not self.para:
                self.para = self.doc.add_paragraph()
            try:
                header, b64_data = data_uri.split(',', 1)
                img_bytes = base64.b64decode(b64_data)
                img_stream = io.BytesIO(img_bytes)
                run = self.para.add_run()
                run.add_picture(img_stream, width=Inches(5.5))
            except Exception:
                self.para.add_run('[图片]')

        def _get_para_obj(self):
            if self.para is None:
                self.para = self.doc.add_paragraph()
            return self.para

        def _apply_alignment(self, style_str):
            if not self.para or not style_str:
                return
            style_str = style_str.lower()
            if 'left' in style_str:
                self.para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif 'center' in style_str:
                self.para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'right' in style_str:
                self.para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        def handle_starttag(self, tag, attrs):
            attrs_dict = dict(attrs)
            style = attrs_dict.get('style', '')

            if tag in ('h1', 'h2', 'h3'):
                self.para = self.doc.add_heading(level=int(tag[1]))
                self._apply_alignment(style)
            elif tag == 'p':
                self.para = self.doc.add_paragraph()
                self._apply_alignment(style)
            elif tag == 'br':
                if self.para:
                    self.para.add_run('\n')
            elif tag == 'table':
                self.in_table = True
                self.table_data = []
            elif tag == 'tr' and self.in_table:
                self.current_row = []
            elif tag in ('td', 'th') and self.in_table:
                self.in_td = True
                self.td_buffer = []
                self.td_formats = []
            elif tag == 'ul':
                self.list_level += 1
            elif tag == 'ol':
                self.list_level += 1
                self.list_counter.append(1)
            elif tag == 'li':
                self.in_li = True
                self.para = self.doc.add_paragraph(style='List Bullet' if self.list_type != 'ol' else 'List Number')
            elif tag == 'blockquote':
                self.para = self.doc.add_paragraph()
                self.para.paragraph_format.left_indent = Inches(0.5)
            elif tag in ('strong', 'b'):
                self.current_formats['bold'] = True
            elif tag in ('em', 'i'):
                self.current_formats['italic'] = True
            elif tag == 'u':
                self.current_formats['underline'] = True
            elif tag == 's':
                self.current_formats['strikethrough'] = True
            elif tag == 'code':
                self.current_formats['code'] = True
            elif tag == 'mark':
                self.current_formats['highlight'] = style_to_color(style, 'background-color') or '#ffff00'
            elif tag == 'span':
                color = style_to_color(style, 'color')
                if color:
                    self.current_formats['color'] = color
                bg = style_to_color(style, 'background-color')
                if bg:
                    self.current_formats['highlight'] = bg
                fs = style_to_font_size(style)
                if fs:
                    self.current_formats['font_size'] = fs
                ff = style_to_font_family(style)
                if ff:
                    self.current_formats['font_family'] = ff
            elif tag == 'img':
                src = attrs_dict.get('src', '')
                if src.startswith('data:image'):
                    self._add_image(src)
                else:
                    if not self.para:
                        self.para = self.doc.add_paragraph()
                    self.para.add_run(f'[图片: {src}]')
            elif tag == 'a':
                self.current_formats['link_url'] = attrs_dict.get('href', '')
            elif tag == 'hr':
                self.para = self.doc.add_paragraph()
                self.para.paragraph_format.space_before = Pt(6)
                self.para.paragraph_format.space_after = Pt(6)
                pPr = self.para._element.pPr
                if pPr is None:
                    pPr = self.para._element.makeelement(qn('w:pPr'), {})
                    self.para._element.insert(0, pPr)
                pBdr = pPr.makeelement(qn('w:pBdr'), {})
                bottom = pBdr.makeelement(qn('w:bottom'), {
                    qn('w:val'): 'single',
                    qn('w:sz'): '6',
                    qn('w:space'): '1',
                    qn('w:color'): 'auto',
                })
                pBdr.append(bottom)
                pPr.append(pBdr)
                self.para = None

        def handle_endtag(self, tag):
            if tag in ('h1', 'h2', 'h3', 'p', 'blockquote'):
                self.para = None
            elif tag == 'table':
                self.in_table = False
                if hasattr(self, 'table_data') and self.table_data:
                    try:
                        rows = len(self.table_data)
                        cols = max(len(r) for r in self.table_data) if self.table_data else 0
                        if rows > 0 and cols > 0:
                            table = self.doc.add_table(rows=rows, cols=cols, style='Table Grid')
                            for ri, row_data in enumerate(self.table_data):
                                for ci, cell_data in enumerate(row_data):
                                    if ci < cols:
                                        cell = table.cell(ri, ci)
                                        cell.text = cell_data
                    except Exception:
                        pass
                self.table_data = []
                self.para = None
            elif tag == 'tr' and self.in_table:
                if hasattr(self, 'current_row'):
                    self.table_data.append(self.current_row)
            elif tag in ('td', 'th') and self.in_table:
                self.in_td = False
                if hasattr(self, 'current_row'):
                    self.current_row.append(''.join(self.td_buffer) if hasattr(self, 'td_buffer') else '')
            elif tag == 'ul':
                self.list_level -= 1
            elif tag == 'ol':
                self.list_level -= 1
                if self.list_counter:
                    self.list_counter.pop()
            elif tag == 'li':
                self.in_li = False
                self.para = None
            elif tag in ('strong', 'b'):
                self.current_formats['bold'] = False
            elif tag in ('em', 'i'):
                self.current_formats['italic'] = False
            elif tag == 'u':
                self.current_formats['underline'] = False
            elif tag == 's':
                self.current_formats['strikethrough'] = False
            elif tag == 'code':
                self.current_formats['code'] = False
            elif tag == 'mark':
                self.current_formats['highlight'] = None
            elif tag == 'span':
                self.current_formats['color'] = None
                self.current_formats['highlight'] = None
                self.current_formats['font_size'] = None
                self.current_formats['font_family'] = None
            elif tag == 'a':
                self.current_formats['link_url'] = None

        def handle_data(self, data):
            if self.in_td:
                self.td_buffer.append(data)
                return
            if not data.strip():
                # Preserve single spaces between inline elements
                if data and not data.startswith('\n') and not data.endswith('\n'):
                    self._add_run(data, self.current_formats.copy())
                return
            self._add_run(data, self.current_formats.copy())

    def style_to_color(style_str, key):
        """Extract color from a CSS style string."""
        import re
        for part in style_str.split(';'):
            part = part.strip()
            if ':' in part:
                k, v = part.split(':', 1)
                if k.strip().lower() == key:
                    color = v.strip()
                    if color.startswith('#'):
                        return color
                    # Handle rgb() and named colors
                    if color.startswith('rgb'):
                        return color
        return None

    def style_to_font_size(style_str):
        import re
        for part in style_str.split(';'):
            part = part.strip()
            if ':' in part:
                k, v = part.split(':', 1)
                if k.strip().lower() == 'font-size':
                    return v.strip()
        return None

    def style_to_font_family(style_str):
        for part in style_str.split(';'):
            part = part.strip()
            if ':' in part:
                k, v = part.split(':', 1)
                if k.strip().lower() == 'font-family':
                    return v.strip().split(',')[0].strip().strip("'\"")
        return None

    parser = DocxHTMLParser()
    try:
        parser.feed(html)
    except Exception:
        # Fallback: strip tags and add plain text
        text = re.sub(r'<[^>]+>', '', html)
        text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
        text = text.replace('&nbsp;', ' ').replace('&quot;', '"')
        for line in text.split('\n'):
            line = line.strip()
            if line:
                doc.add_paragraph(line)


def generate_xlsx(content: str, name: str) -> dict:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    rows_data = []
    table_match = re.search(r'<table[^>]*>(.+?)</table>', content, re.DOTALL)
    if table_match:
        for tr_match in re.finditer(r'<tr[^>]*>(.+?)</tr>', table_match.group(1), re.DOTALL):
            cells = []
            for td_match in re.finditer(r'<t[dh][^>]*>(.*?)</t[dh]>', tr_match.group(1), re.DOTALL):
                cell_text = re.sub(r'<[^>]+>', '', td_match.group(1)).strip()
                cells.append(cell_text)
            if cells:
                rows_data.append(cells)
    else:
        for line in content.split('\n'):
            line = line.strip()
            if line:
                rows_data.append([line])

    for ri, row in enumerate(rows_data, 1):
        for ci, val in enumerate(row, 1):
            cell = ws.cell(row=ri, column=ci)
            cell.value = val
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', vertical='center')

    tmp = tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False)
    wb.save(tmp.name)
    wb.close()
    with open(tmp.name, 'rb') as f:
        b64 = base64.b64encode(f.read()).decode('ascii')
    os.unlink(tmp.name)
    return {"success": True, "data_uri": f"data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}"}
